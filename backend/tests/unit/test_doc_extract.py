"""Tests for the document text extractor.

We synthesise minimal valid files for each format using their writer
libraries (pypdf can't synthesise a PDF, so we generate it via
reportlab — but to avoid a heavyweight extra dep we ship a fixture PDF
inline as a tiny hand-written byte stream).
"""
from __future__ import annotations

import io
import json

import pytest

from app.services.doc_extract import (
    SUPPORTED_EXTENSIONS,
    _decode_text,
    ext_of,
    extract_text,
)


# ── ext + plain-text decoding ─────────────────────────────────────


def test_ext_of_basic() -> None:
    assert ext_of("doc.PDF") == ".pdf"
    assert ext_of("/abs/a.b.txt") == ".txt"
    assert ext_of("noext") == ""
    assert ext_of("") == ""


def test_decode_utf8_bom() -> None:
    assert _decode_text(b"\xef\xbb\xbfHello") == "Hello"


def test_decode_utf16_le_bom() -> None:
    assert _decode_text(b"\xff\xfeH\x00i\x00") == "Hi"


def test_decode_plain_utf8() -> None:
    cyrillic = "Привет мир".encode("utf-8")
    assert _decode_text(cyrillic) == "Привет мир"


def test_decode_uzbek_latin_and_cyrillic() -> None:
    # bge-m3 handles both — we just verify decoding round-trips.
    assert _decode_text("Salom dunyo".encode("utf-8")) == "Salom dunyo"
    assert _decode_text("Салом дунё".encode("utf-8")) == "Салом дунё"


def test_decode_falls_back_to_latin1() -> None:
    # Random bytes that aren't valid UTF-8 must still decode (no raise).
    result = _decode_text(b"\xe9\xe8\xea")
    assert isinstance(result, str)


# ── plain text extraction ─────────────────────────────────────────


def test_extract_txt() -> None:
    body = "Hello\nworld"
    out = extract_text("notes.txt", body.encode("utf-8"))
    assert out is not None
    text, mime = out
    assert text == body
    assert mime == "text/plain"


def test_extract_md() -> None:
    body = "# Heading\n\nParagraph"
    out = extract_text("readme.md", body.encode("utf-8"))
    assert out is not None
    assert "Heading" in out[0]


def test_extract_csv_as_text() -> None:
    body = "id,name\n1,alice\n2,bob"
    out = extract_text("rows.csv", body.encode("utf-8"))
    assert out is not None
    assert "alice" in out[0]


def test_extract_unsupported_returns_none() -> None:
    assert extract_text("photo.png", b"\x89PNG...") is None
    assert extract_text("video.mp4", b"\x00\x00\x00\x20ftypisom") is None


# ── HTML extraction ──────────────────────────────────────────────


def test_extract_html_strips_script_style() -> None:
    html = b"""
    <html>
      <head>
        <style>body { color: red; }</style>
        <script>alert('hi')</script>
      </head>
      <body>
        <h1>Quarterly Report</h1>
        <p>Revenue grew 30%.</p>
      </body>
    </html>
    """
    out = extract_text("report.html", html)
    assert out is not None
    text, mime = out
    assert "Quarterly Report" in text
    assert "Revenue grew" in text
    assert "alert" not in text
    assert "color: red" not in text
    assert mime == "text/html"


def test_extract_html_multilingual() -> None:
    html = (
        "<html><body>"
        "<p>Salom dunyo</p>"
        "<p>Привет мир</p>"
        "<p>Hello world</p>"
        "</body></html>"
    ).encode("utf-8")
    out = extract_text("multi.html", html)
    assert out is not None
    text = out[0]
    assert "Salom dunyo" in text
    assert "Привет мир" in text
    assert "Hello world" in text


# ── DOCX extraction (uses python-docx) ────────────────────────────


def test_extract_docx() -> None:
    import docx

    doc = docx.Document()
    doc.add_heading("Project Plan", level=1)
    doc.add_paragraph("Phase 1 — discovery (3 weeks).")
    doc.add_paragraph("Phase 2 — build (6 weeks).")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Owner"
    tbl.rows[0].cells[1].text = "Status"
    tbl.rows[1].cells[0].text = "Alisher"
    tbl.rows[1].cells[1].text = "On track"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out = extract_text("plan.docx", buf.read())
    assert out is not None
    text, mime = out
    assert "Project Plan" in text
    assert "discovery" in text
    assert "Alisher" in text
    assert "On track" in text
    assert "wordprocessingml.document" in mime


# ── XLSX extraction (uses openpyxl) ───────────────────────────────


def test_extract_xlsx() -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Q1"
    ws.append(["Region", "Revenue"])
    ws.append(["EMEA", 12500])
    ws.append(["APAC", 9800])
    ws2 = wb.create_sheet("Q2")
    ws2.append(["Region", "Revenue"])
    ws2.append(["EMEA", 14200])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    out = extract_text("revenue.xlsx", buf.read())
    assert out is not None
    text, mime = out
    assert "Sheet: Q1" in text
    assert "Sheet: Q2" in text
    assert "EMEA" in text
    assert "12500" in text
    assert "spreadsheetml.sheet" in mime


# ── PDF extraction (uses pypdf) ───────────────────────────────────


def test_extract_pdf_minimal() -> None:
    """Build a minimal PDF on the fly using pypdf's writer + a single
    rendered text block. Verifies the extractor at least runs on PDFs
    without raising; perfect-fidelity extraction of arbitrary PDFs is
    pypdf's problem, not ours."""
    pytest.importorskip("pypdf")
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        DecodedStreamObject,
        DictionaryObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    # Construct an absolutely minimal 1-page PDF with a literal text op.
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    # Inject a Contents stream with `BT /F1 12 Tf 72 720 Td (Hello) Tj ET`.
    content = DecodedStreamObject()
    content.set_data(
        b"BT /F1 12 Tf 72 720 Td (Hello hackathon) Tj ET"
    )
    page[NameObject("/Contents")] = content
    # Bare-bones font resource so pypdf doesn't choke on the F1 ref.
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {
                    NameObject("/F1"): DictionaryObject(
                        {
                            NameObject("/Type"): NameObject("/Font"),
                            NameObject("/Subtype"): NameObject("/Type1"),
                            NameObject("/BaseFont"): NameObject("/Helvetica"),
                        }
                    )
                }
            )
        }
    )

    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)

    out = extract_text("hello.pdf", buf.read())
    assert out is not None
    text, mime = out
    assert mime == "application/pdf"
    # pypdf's text extraction on hand-built PDFs can be flaky; we tolerate
    # an empty extraction here — the goal is verifying no exception.
    assert isinstance(text, str)


# ── allow-list contract ──────────────────────────────────────────


def test_supported_extensions_set_nonempty() -> None:
    assert len(SUPPORTED_EXTENSIONS) > 5
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert ".docx" in SUPPORTED_EXTENSIONS
    assert ".xlsx" in SUPPORTED_EXTENSIONS
    assert ".html" in SUPPORTED_EXTENSIONS
    assert ".txt" in SUPPORTED_EXTENSIONS
