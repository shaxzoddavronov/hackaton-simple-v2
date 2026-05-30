"""Extract plain text from common document formats.

Used by the Phase 14 harvester to turn an arbitrary file's bytes into
something we can chunk + embed. The extractor dispatches on the file
extension (lower-cased) and returns ``(text, mime_type)``; ``None`` is
returned for unsupported formats so callers can skip without
exceptions.

Supported:
  * ``.pdf``                — pypdf for text-PDFs; pdf2image +
    Tesseract OCR fallback when pypdf returns an empty page (scanned
    document, image-only).
  * ``.docx``               — python-docx
  * ``.xlsx``, ``.xlsm``    — openpyxl (read each sheet as CSV-ish text)
  * ``.html``, ``.htm``     — BeautifulSoup, strip <script>/<style>
  * ``.txt``, ``.md``, ``.csv``, ``.tsv``, ``.json``, ``.log``
                            — direct UTF-8 decode (with BOM detection)
  * ``.mp3``, ``.mp4``, ``.m4a``, ``.wav``, ``.webm``, ``.ogg``,
    ``.opus``, ``.mpeg``, ``.mpga``
                            — faster-whisper speech-to-text. Language
    is auto-detected (Uzbek / Russian / English all work). Requires
    ``ffmpeg`` on PATH for non-WAV decoding — see
    ``infra/README.md`` for install notes.
  * ``.png``, ``.jpg``/``.jpeg``, ``.tiff``/``.tif``, ``.bmp``,
    ``.webp``               — Tesseract OCR via ``pytesseract``.
    Multilingual: uzb+rus+eng tessdata picked up automatically when
    the language packs are installed. See ``infra/README.md`` for
    the Tesseract install path.

The hard cap on extracted text per file is :data:`MAX_TEXT_BYTES`
(default 1 MB). Anything beyond is truncated — embedding a 50-page
PDF as one chunk would blow the model's context anyway, and the
chunker downstream slices the text into 1200-char windows.

Embedding model is bge-m3 which natively supports 100+ languages
including Uzbek, Russian and English — no language-specific
preprocessing needed.
"""
from __future__ import annotations

import io
import logging
import os
import tempfile
import threading
from typing import Optional

log = logging.getLogger(__name__)

MAX_TEXT_BYTES = 1 * 1024 * 1024  # 1 MB extracted text cap per file


# Extensions we treat as plain UTF-8 text on disk. We sniff the BOM
# (UTF-8 / UTF-16 LE / UTF-16 BE) so files saved by Windows tools
# decode cleanly.
_PLAIN_TEXT_EXTS = {
    ".txt", ".md", ".csv", ".tsv", ".log", ".json", ".jsonl", ".ndjson",
    ".yaml", ".yml", ".sql", ".ini", ".cfg", ".conf",
}

# Audio/video extensions transcribed with Whisper. ``.mp4`` / ``.webm``
# may contain video — Whisper (via ffmpeg) just demuxes the audio
# track. Anything ffmpeg can decode is fair game; this set is the
# common subset Whisper documents as supported.
_AUDIO_EXTS = {
    ".mp3", ".m4a", ".wav", ".webm", ".ogg", ".opus",
    ".mp4", ".mpeg", ".mpga",
}

# Raster image extensions OCRed via Tesseract. The harvester pulls
# these out of folders / SMB shares / cloud drives like any other
# file; downstream the chunker stores the recognised text and bge-m3
# embeds it. Multilingual is one ``-l uzb+rus+eng`` flag away (see
# OCR_LANGS env).
_IMAGE_EXTS = {
    ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp",
}

# Extensions where we will produce an extraction. Used by the API
# layer to allow-list uploads on the source side.
SUPPORTED_EXTENSIONS = (
    _PLAIN_TEXT_EXTS
    | {".pdf", ".docx", ".xlsx", ".xlsm", ".html", ".htm"}
    | _AUDIO_EXTS
    | _IMAGE_EXTS
)


# Tesseract language string. Default to uzb+rus+eng so all three
# project-target languages work without per-source configuration.
# Override via env if you've installed extra tessdata packs.
OCR_LANGS = os.environ.get("OCR_LANGS", "uzb+rus+eng")
# ``OCR_PDF_DPI`` controls the render resolution for the scanned-PDF
# fallback. 200 DPI is a sweet spot between speed and accuracy on
# typical receipts / forms; raise to 300 for fine print.
OCR_PDF_DPI = int(os.environ.get("OCR_PDF_DPI", "200"))


# MIME types for the audio formats we transcribe. ``.mp4`` is reported
# as ``video/mp4`` because that's its canonical IANA type even though
# we only look at the audio track.
_AUDIO_MIME = {
    ".mp3":  "audio/mpeg",
    ".mpga": "audio/mpeg",
    ".mpeg": "audio/mpeg",
    ".m4a":  "audio/mp4",
    ".wav":  "audio/wav",
    ".webm": "audio/webm",
    ".ogg":  "audio/ogg",
    ".opus": "audio/opus",
    ".mp4":  "video/mp4",
}


def _mime_for_audio(ext: str) -> str:
    return _AUDIO_MIME.get(ext, "application/octet-stream")


def ext_of(filename: str) -> str:
    """Return the lower-cased extension including the dot. ``''`` if
    the file has no extension."""
    _, ext = os.path.splitext(filename or "")
    return ext.lower()


def extract_text(filename: str, data: bytes) -> Optional[tuple[str, str]]:
    """Return ``(text, mime_type)`` or ``None`` for unsupported files.

    Errors during extraction (corrupt PDF, password-protected DOCX,
    etc.) bubble up as exceptions — the caller decides whether to mark
    the file as failed or retry. We don't silently swallow because
    silent failures would let the index drift from the source.
    """
    ext = ext_of(filename)
    if ext == ".pdf":
        return _extract_pdf(data), "application/pdf"
    if ext == ".docx":
        return _extract_docx(data), (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )
    if ext in (".xlsx", ".xlsm"):
        return _extract_xlsx(data), (
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        )
    if ext in (".html", ".htm"):
        return _extract_html(data), "text/html"
    if ext in _PLAIN_TEXT_EXTS:
        return _decode_text(data), "text/plain"
    if ext in _AUDIO_EXTS:
        return _extract_audio(filename, data), _mime_for_audio(ext)
    if ext in _IMAGE_EXTS:
        return _extract_image(data), f"image/{ext.lstrip('.')}"
    return None


# ── per-format helpers ────────────────────────────────────────────


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF.

    Two-pass strategy (Phase 20):
      1. Try ``pypdf`` — fast, no system deps, works for any PDF that
         was actually authored from text (Office exports, LaTeX, web
         to PDF, etc.).
      2. If pypdf returns an empty string — typical for scanned
         documents where every page is an image — fall back to
         rendering each page with ``pypdfium2`` and running Tesseract
         OCR via ``pytesseract``. This is slow (~1–3 s per page on
         CPU) but recovers text from scans the user expects to be
         searchable.

    The OCR fallback only triggers when the entire text extraction
    came back empty, so text-PDFs don't pay the OCR cost. A
    half-image-half-text PDF (e.g. an Office doc with a scanned
    appendix) is still served from the text path; the scanned pages
    contribute nothing but the rest of the document is captured.
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    total = 0
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception as e:
            # Per-page failures shouldn't kill the whole doc; pypdf
            # occasionally trips on malformed fonts / encrypted pages.
            log.warning("pdf page %d extraction failed: %s", i, e)
            continue
        if t:
            parts.append(t)
            total += len(t.encode("utf-8"))
            if total > MAX_TEXT_BYTES:
                break
    text = _join_paragraphs(parts)

    # Empty pypdf output → almost certainly a scan. Drop into OCR.
    if not text.strip():
        ocr_text = _extract_pdf_ocr(data)
        if ocr_text:
            return _cap(ocr_text)
    return _cap(text)


def _extract_pdf_ocr(data: bytes) -> str:
    """Render every page of ``data`` and OCR it with Tesseract.

    Uses ``pypdfium2`` (statically-linked PDFium fork) so there's no
    poppler dependency on the host. Per-page text is joined with
    paragraph separators. Bails the moment the running total crosses
    :data:`MAX_TEXT_BYTES` — embeds + chunks are already capped, so
    a 500-page archive doesn't sit in RAM forever.
    """
    try:
        import pypdfium2 as pdfium
    except ImportError:
        log.warning(
            "scanned PDF detected but pypdfium2 is not installed — "
            "install it to enable OCR fallback"
        )
        return ""

    parts: list[str] = []
    total = 0
    try:
        doc = pdfium.PdfDocument(data)
    except Exception as e:
        log.warning("pypdfium2 failed to open PDF: %s", e)
        return ""
    try:
        scale = OCR_PDF_DPI / 72.0  # PDF native is 72 DPI
        for i in range(len(doc)):
            try:
                page = doc[i]
                # render_to returns a PIL Image when ``PIL`` is asked
                # for. ``render`` (no ``_to``) is the newer API; we
                # use the version-agnostic shape.
                pil_img = page.render(scale=scale).to_pil()
            except Exception as e:
                log.warning("pdf ocr: render page %d failed: %s", i, e)
                continue
            try:
                page_text = _ocr_image(pil_img)
            except Exception as e:
                log.warning("pdf ocr: tesseract page %d failed: %s", i, e)
                continue
            if page_text:
                parts.append(page_text)
                total += len(page_text.encode("utf-8"))
                if total > MAX_TEXT_BYTES:
                    break
    finally:
        try:
            doc.close()
        except Exception:
            pass
    return _join_paragraphs(parts)


def _extract_image(data: bytes) -> str:
    """OCR a single image. Returns empty string when Tesseract can't
    extract anything (blurry / no text / wrong language pack)."""
    try:
        from PIL import Image
    except ImportError:
        log.warning(
            "image OCR requested but Pillow is not installed — "
            "install pillow + pytesseract to enable image extraction"
        )
        return ""
    try:
        img = Image.open(io.BytesIO(data))
    except Exception as e:
        log.warning("image OCR: PIL failed to open: %s", e)
        return ""
    try:
        text = _ocr_image(img)
    except Exception as e:
        log.warning("image OCR: tesseract failed: %s", e)
        return ""
    return _cap(text)


def _ocr_image(pil_image) -> str:
    """Single dispatch into Tesseract. Wraps the FileNotFoundError
    case so a missing binary surfaces with a clear install hint."""
    try:
        import pytesseract
    except ImportError as e:
        raise RuntimeError(
            "pytesseract not installed — pip install pytesseract"
        ) from e
    try:
        return pytesseract.image_to_string(pil_image, lang=OCR_LANGS) or ""
    except FileNotFoundError as e:
        # Same pattern as the Whisper / ffmpeg path: turn the cryptic
        # binary-not-found into something actionable.
        raise RuntimeError(
            "tesseract binary not found on PATH — install Tesseract OCR "
            "and the language packs uzb+rus+eng. "
            "Linux: `apt-get install tesseract-ocr tesseract-ocr-uzb "
            "tesseract-ocr-rus`. "
            "Windows: install from "
            "https://github.com/UB-Mannheim/tesseract/wiki and add the "
            "install folder to PATH."
        ) from e


def _extract_docx(data: bytes) -> str:
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    # Body paragraphs.
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # Tables — flatten each row as a tab-separated line so the chunker
    # keeps cells together.
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [cell.text or "" for cell in row.cells]
            line = "\t".join(cells).strip()
            if line:
                parts.append(line)
    return _cap(_join_paragraphs(parts))


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        parts.append(f"## Sheet: {sheet_name}")
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells).strip()
            if line:
                parts.append(line)
    return _cap(_join_paragraphs(parts))


def _extract_html(data: bytes) -> str:
    from bs4 import BeautifulSoup

    # html.parser is stdlib; lxml would be faster but we avoid the
    # extra wheel.
    soup = BeautifulSoup(_decode_text(data), "html.parser")
    # Drop script/style and other noise that contributes no semantics.
    for tag in soup(["script", "style", "noscript", "iframe"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    # Collapse runs of blank lines.
    lines = [ln.strip() for ln in text.splitlines()]
    return _cap("\n".join(ln for ln in lines if ln))


def _decode_text(data: bytes) -> str:
    """Decode bytes as text with simple BOM detection."""
    if data.startswith(b"\xef\xbb\xbf"):
        return data[3:].decode("utf-8", errors="replace")
    if data.startswith(b"\xff\xfe"):
        return data[2:].decode("utf-16-le", errors="replace")
    if data.startswith(b"\xfe\xff"):
        return data[2:].decode("utf-16-be", errors="replace")
    # Try UTF-8 first; fall back to latin-1 which can't fail.
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


# ── helpers ────────────────────────────────────────────────────────


def _join_paragraphs(parts: list[str]) -> str:
    out: list[str] = []
    for p in parts:
        p = p.strip()
        if p:
            out.append(p)
    return "\n\n".join(out)


def _cap(text: str) -> str:
    if len(text.encode("utf-8")) <= MAX_TEXT_BYTES:
        return text
    # Cut at a UTF-8 boundary to avoid lone surrogates.
    encoded = text.encode("utf-8")[:MAX_TEXT_BYTES]
    return encoded.decode("utf-8", errors="ignore") + "\n[...truncated]"


# ── Whisper audio transcription ──────────────────────────────────
#
# Loaded lazily on first transcription — most QueryMind processes
# (API server, scheduler, vLLM client) never see audio, so paying the
# ~140 MB model download + 200 MB resident memory at import time
# would be wasteful. The double-checked lock around the singleton is
# the canonical pattern for "import-once, share-across-threads" in
# Celery prefork workers where workers fork the parent and inherit
# its (post-init) module state.

_whisper_model = None
_whisper_lock = threading.Lock()


def _get_whisper():
    """Return the process-wide ``faster_whisper.WhisperModel`` singleton.

    Size + compute backend are configurable via env so a dev box with a
    GPU can opt into ``WHISPER_DEVICE=cuda`` / ``WHISPER_COMPUTE_TYPE=float16``
    without a code change. Defaults target CPU-only QueryMind dev boxes.
    """
    global _whisper_model
    if _whisper_model is None:
        with _whisper_lock:
            if _whisper_model is None:
                from faster_whisper import WhisperModel

                _whisper_model = WhisperModel(
                    os.environ.get("WHISPER_MODEL_SIZE", "base"),
                    device=os.environ.get("WHISPER_DEVICE", "cpu"),
                    compute_type=os.environ.get(
                        "WHISPER_COMPUTE_TYPE", "int8"
                    ),
                )
    return _whisper_model


def _extract_audio(filename: str, data: bytes) -> str:
    """Transcribe an audio/video file to plain text via faster-whisper.

    The bytes are written to a NamedTemporaryFile because faster-whisper
    (and the CTranslate2 backend underneath) wants a file path — passing
    an in-memory BytesIO works for some codecs but not all, and the
    temp-file path is the documented happy path.

    Language is auto-detected (``language=None``) so Uzbek, Russian,
    and English clips all transcribe without extra configuration. On
    a corrupted / silent clip Whisper returns zero segments and we
    return an empty string rather than raising — the harvester logs
    the empty result and moves on to the next file.
    """
    ext = ext_of(filename) or ".audio"
    model = _get_whisper()

    # delete=False because on Windows the WhisperModel reopens the path
    # internally; if the tempfile is still open under us, the second
    # open fails with PermissionError. We delete in `finally` instead.
    tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
    try:
        tmp.write(data)
        tmp.flush()
        tmp.close()

        try:
            segments, _info = model.transcribe(tmp.name, language=None)
        except FileNotFoundError as e:
            # ffmpeg missing on PATH is the most common cause —
            # faster-whisper invokes ``ffmpeg`` for non-WAV decoding and
            # Python raises FileNotFoundError when the binary is missing.
            if "ffmpeg" in str(e).lower():
                raise RuntimeError(
                    "ffmpeg not found on PATH — install ffmpeg and "
                    "restart the Celery worker. Windows: "
                    "`winget install Gyan.FFmpeg`. Linux: "
                    "`apt-get install ffmpeg`."
                ) from e
            raise

        parts: list[str] = []
        total = 0
        for seg in segments:
            text = (getattr(seg, "text", "") or "").strip()
            if not text:
                continue
            parts.append(text)
            total += len(text.encode("utf-8")) + 1  # +1 for the space
            if total > MAX_TEXT_BYTES:
                # Stop pulling more segments — _cap() will trim cleanly
                # to a UTF-8 boundary below. A 1-hour meeting transcript
                # would otherwise sit at ~50 KB which is fine; this guard
                # is the belt to _cap()'s suspenders.
                break

        transcript = " ".join(parts).strip()
        return _cap(transcript)
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            # Best-effort cleanup; the OS temp sweeper will reclaim it.
            pass
